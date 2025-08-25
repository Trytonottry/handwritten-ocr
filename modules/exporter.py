# modules/exporter.py
from docx import Document
import os

def save_as_txt(pages_data, output_path):
    """Сохраняет результат в TXT"""
    with open(output_path, "w", encoding="utf-8") as f:
        for page_num, content in enumerate(pages_data, 1):
            f.write(f"--- Страница {page_num} ---\n")
            for item in content:
                f.write(f"{item['text']} (уверенность: {item['confidence']:.2f})\n")
            f.write("\n")
    print(f"Текст сохранён в {output_path}")

def save_as_docx(pages_data, output_path):
    """Сохраняет результат в DOCX"""
    doc = Document()
    doc.add_heading('Распознанный текст (рукописный)', 0)

    for page_num, content in enumerate(pages_data, 1):
        doc.add_heading(f'Страница {page_num}', level=1)
        for item in content:
            p = doc.add_paragraph()
            p.add_run(item['text']).bold = item['confidence'] < 0.5
            p.add_run(f"  [уверенность: {item['confidence']:.2f}]").italic = True
    doc.save(output_path)
    print(f"Документ сохранён в {output_path}")